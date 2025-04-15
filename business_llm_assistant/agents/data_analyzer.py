"""
Data Analyzer Agent Module

This module implements the Data Analyzer agent, which specializes in analyzing
business data, processing documents, and extracting insights from various data sources.
"""

import os
import json
import logging
import time
from typing import Dict, List, Any, Optional, Tuple, Union
import re
from datetime import datetime
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pathlib import Path
import base64
from io import BytesIO
import tempfile

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import local modules
try:
    from ..core.coordinator import Agent, Message
    from ..core.model_loader import ModelLoader
except ImportError:
    # Adjust import path if module is run directly
    import sys
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from core.coordinator import Agent, Message
    from core.model_loader import ModelLoader

# Try to import document processing libraries
try:
    import PyPDF2
    from PIL import Image
    import pytesseract
    PDF_SUPPORT = True
except ImportError:
    logger.warning("PDF/Image processing libraries not available. Install with pip install PyPDF2 pillow pytesseract")
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    logger.warning("Word document processing library not available. Install with pip install python-docx")
    DOCX_SUPPORT = False

try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    logger.warning("Excel processing library not available. Install with pip install openpyxl")
    EXCEL_SUPPORT = False


class DataAnalyzerAgent(Agent):
    """
    Agent specializing in analyzing business data and documents
    """
    
    def __init__(self, name: str, model_name: str = "DeepCoder-14B-Preview-GGUF", 
                 model_type: str = "gguf", **kwargs):
        """
        Initialize the data analyzer agent

        Args:
            name: Agent name
            model_name: Name of model to use
            model_type: Type of model (gguf or hf)
            **kwargs: Additional configuration options
        """
        super().__init__(name, model_name, model_type)
        
        # Set up system message for data analysis tasks
        self.system_message = """You are a sophisticated data analysis assistant specializing in business intelligence,
data processing, and document analysis. Your responsibilities include:

1. Analyzing numerical and textual business data to extract insights
2. Processing various document formats (PDF, Word, Excel, Text) to extract information
3. Creating data visualizations and reporting key metrics
4. Identifying trends and patterns in business data
5. Providing actionable recommendations based on data analysis

You should be precise, factual, and data-driven in your analysis.
When working with numerical data, include specific metrics and statistics.
Always verify the reliability and source of data before drawing conclusions.
"""
        
        # Set up template for LLM prompting
        self.prompt_template = """
{system_message}

# Context:
{context}

# Conversation History:
{conversation_history}

# Current Request:
{current_message}

# Task:
{specific_task}

Respond with a comprehensive data analysis that addresses the request.
If code is needed to process or visualize data, include it.
"""
        
        # Setup document storage path
        self.document_dir = os.path.join(os.getcwd(), "data", "documents")
        os.makedirs(self.document_dir, exist_ok=True)
        
        # Setup vector database path
        self.vector_db_dir = os.path.join(os.getcwd(), "data", "vector_db")
        os.makedirs(self.vector_db_dir, exist_ok=True)
        
        # Configure from kwargs
        self.max_history = kwargs.get("max_history", 10)
        self.context_length = kwargs.get("context_length", 4096)
        self.temperature = kwargs.get("temperature", 0.7)
        self.top_p = kwargs.get("top_p", 0.9)
        
        # Check for vision model capability
        self.vision_model_name = kwargs.get("vision_model_name", "Kimi-VL-A3B")
        self.use_vision_model = kwargs.get("use_vision_model", True)
        self.vision_model = None
        self.vision_processor = None
    
    def initialize(self, model_loader: ModelLoader) -> bool:
        """
        Initialize the agent with the appropriate model

        Args:
            model_loader: ModelLoader instance to load the model

        Returns:
            True if successful, False otherwise
        """
        try:
            # Load the primary model based on model type
            if self.model_type.lower() == "gguf":
                # Load GGUF model
                self.model = model_loader.load_gguf_model(
                    model_path=self.model_name,
                    agent_name=self.name,
                    n_ctx=self.context_length,
                    model_kwargs={"n_threads": 4}  # Use 4 threads for better CPU inference
                )
                logger.info(f"Loaded GGUF model for {self.name}")
            else:
                # Load Hugging Face model
                self.model, self.tokenizer = model_loader.load_hf_model(
                    model_name=self.model_name,
                    agent_name=self.name,
                    quantization="4bit" 
                )
                logger.info(f"Loaded HF model for {self.name}")
            
            # Try to load vision model if enabled
            if self.use_vision_model:
                try:
                    self.vision_model, self.vision_processor = model_loader.load_hf_model(
                        model_name=self.vision_model_name,
                        agent_name=f"{self.name}_vision",
                        model_type="multimodal"
                    )
                    logger.info(f"Loaded vision model {self.vision_model_name}")
                except Exception as e:
                    logger.warning(f"Failed to load vision model: {str(e)}")
                    self.use_vision_model = False
            
            return True
                
        except Exception as e:
            logger.error(f"Error initializing {self.name}: {str(e)}")
            return False
    
    def process_message(self, message: Message) -> Message:
        """
        Process a message and generate a response

        Args:
            message: Message to process

        Returns:
            Response message
        """
        # Update conversation history
        self.update_history(message)
        
        # Check if it's a data analysis request or document processing request
        content = message.content.lower()
        
        # Extract context based on the request
        context = ""
        files_to_analyze = []
        
        # Look for file references in the message
        if re.search(r'(analyze|examine|review|process|read|extract from)\s+([a-zA-Z0-9_\-\./]+\.(pdf|docx?|xlsx?|csv|txt|png|jpg|jpeg))', message.content, re.IGNORECASE):
            # Extract filenames
            matches = re.findall(r'([a-zA-Z0-9_\-\./]+\.(pdf|docx?|xlsx?|csv|txt|png|jpg|jpeg))', message.content)
            for match in matches:
                files_to_analyze.append(match[0])
            
            # Process files and add to context
            for filename in files_to_analyze:
                file_content = self._process_document(filename)
                if file_content:
                    context += f"Content from file '{filename}':\n{file_content}\n\n"
        
        # Determine specific task based on message content
        specific_task = self._determine_task(message.content)
        
        # Generate prompt
        prompt = self.get_prompt(message, context, specific_task)
        
        try:
            # Generate response
            if self.model_type.lower() == "gguf":
                # Generate with GGUF model
                response_text = self._generate_gguf(prompt)
            else:
                # Generate with HF model
                response_text = self._generate_hf(prompt)
            
            # Check if response contains plots to generate
            plot_code = self._extract_plot_code(response_text)
            if plot_code:
                try:
                    plot_images = self._execute_plot_code(plot_code)
                    if plot_images:
                        response_text = self._add_plot_images_to_response(response_text, plot_images)
                except Exception as e:
                    logger.error(f"Error generating plots: {str(e)}")
                    response_text += f"\n\nNote: There was an error generating the plots: {str(e)}"
            
            # Create response message
            response = Message(
                content=response_text,
                sender=self.name,
                receiver=message.sender,
                message_type="response",
                metadata={
                    "analyzed_files": files_to_analyze
                } if files_to_analyze else {}
            )
            
            # Update history with our response
            self.update_history(response)
            
            return response
            
        except Exception as e:
            logger.error(f"Error processing message in {self.name}: {str(e)}")
            
            # Return error message
            return Message(
                content=f"I encountered an error while analyzing the data: {str(e)}",
                sender=self.name,
                receiver=message.sender,
                message_type="error"
            )
    
    def get_prompt(self, message: Message, context: str = "", specific_task: str = "") -> str:
        """
        Generate prompt from message and conversation history

        Args:
            message: Current message
            context: Additional context from documents
            specific_task: Specific task description

        Returns:
            Complete prompt for the model
        """
        # Format conversation history
        history_text = ""
        for i, hist_msg in enumerate(self.conversation_history[-self.max_history:-1]):
            role = "User" if hist_msg["sender"] == "user" else "Assistant"
            history_text += f"{role}: {hist_msg['content']}\n\n"
        
        # Format prompt
        prompt = self.prompt_template.format(
            system_message=self.system_message,
            context=context,
            conversation_history=history_text,
            current_message=message.content,
            specific_task=specific_task
        )
        
        return prompt
    
    def _determine_task(self, message_content: str) -> str:
        """
        Determine the specific task based on message content

        Args:
            message_content: Message content

        Returns:
            Task description string
        """
        content = message_content.lower()
        
        if any(word in content for word in ["analyze pdf", "extract from pdf", "read pdf"]):
            return "Extract and analyze information from the provided PDF document."
        
        elif any(word in content for word in ["analyze image", "look at image", "what's in this image"]):
            return "Analyze the provided image and extract relevant information."
        
        elif any(word in content for word in ["analyze spreadsheet", "excel data", "csv data"]):
            return "Analyze the provided spreadsheet/CSV data and extract key insights."
        
        elif any(word in content for word in ["visualize", "create chart", "plot", "graph"]):
            return "Create appropriate data visualizations based on the provided data."
        
        elif any(word in content for word in ["trend", "pattern", "correlation"]):
            return "Identify trends, patterns, or correlations in the provided data."
        
        elif any(word in content for word in ["predict", "forecast", "projection"]):
            return "Provide predictive analysis or forecasting based on historical data."
        
        elif any(word in content for word in ["summarize", "summary", "key points"]):
            return "Summarize key findings and insights from the data."
        
        else:
            return "Analyze the provided data or request and provide useful insights."
    
    def _process_document(self, filename: str) -> str:
        """
        Process a document and extract its content

        Args:
            filename: Document filename

        Returns:
            Extracted text content
        """
        # Get full path
        if not os.path.isabs(filename):
            file_path = os.path.join(self.document_dir, filename)
        else:
            file_path = filename
        
        # Check if file exists
        if not os.path.exists(file_path):
            logger.warning(f"File not found: {file_path}")
            return f"ERROR: File '{filename}' not found."
        
        # Process based on file extension
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            # PDF file
            if file_ext == '.pdf' and PDF_SUPPORT:
                return self._extract_text_from_pdf(file_path)
            
            # Word document
            elif file_ext in ['.doc', '.docx'] and DOCX_SUPPORT:
                return self._extract_text_from_docx(file_path)
            
            # Excel spreadsheet
            elif file_ext in ['.xls', '.xlsx'] and EXCEL_SUPPORT:
                return self._extract_text_from_excel(file_path)
            
            # CSV file
            elif file_ext == '.csv':
                return self._extract_text_from_csv(file_path)
            
            # Text file
            elif file_ext == '.txt':
                return self._extract_text_from_txt(file_path)
            
            # Image file
            elif file_ext in ['.png', '.jpg', '.jpeg'] and self.use_vision_model and self.vision_model is not None:
                # Use vision model to analyze the image
                return self._analyze_image(file_path)
            
            # Unsupported file type
            else:
                return f"ERROR: Unsupported file type '{file_ext}'"
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            return f"ERROR: Could not process file '{filename}': {str(e)}"
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text
        """
        if not PDF_SUPPORT:
            return "PDF processing is not available. Please install PyPDF2."
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += f"--- Page {page_num+1} ---\n"
                text += page.extract_text()
                text += "\n\n"
            
            # If no text was extracted, it might be a scanned PDF
            if not text.strip():
                logger.info(f"No text extracted from PDF, trying OCR: {file_path}")
                # Try using OCR if available
                if pytesseract:
                    try:
                        from pdf2image import convert_from_path
                        images = convert_from_path(file_path)
                        for i, image in enumerate(images):
                            text += f"--- Page {i+1} ---\n"
                            text += pytesseract.image_to_string(image)
                            text += "\n\n"
                    except Exception as e:
                        logger.error(f"OCR failed for PDF: {str(e)}")
                        text = "This appears to be a scanned PDF. OCR processing failed."
            
            return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a Word document

        Args:
            file_path: Path to the Word document

        Returns:
            Extracted text
        """
        if not DOCX_SUPPORT:
            return "Word document processing is not available. Please install python-docx."
        
        doc = docx.Document(file_path)
        text = ""
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract tables if any
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text += " | ".join(row_text) + "\n"
        
        return text
    
    def _extract_text_from_excel(self, file_path: str) -> str:
        """
        Extract data from an Excel file

        Args:
            file_path: Path to the Excel file

        Returns:
            Extracted data as text
        """
        if not EXCEL_SUPPORT:
            return "Excel processing is not available. Please install openpyxl."
        
        # Read Excel file
        df_dict = pd.read_excel(file_path, sheet_name=None)
        text = ""
        
        # Process each sheet
        for sheet_name, df in df_dict.items():
            text += f"--- Sheet: {sheet_name} ---\n"
            
            # Get column headers
            headers = df.columns.tolist()
            text += "Columns: " + ", ".join([str(h) for h in headers]) + "\n\n"
            
            # Basic statistics
            text += "Basic Statistics:\n"
            for col in df.columns:
                if pd.api.types.is_numeric_dtype(df[col]):
                    text += f"- {col}:\n"
                    text += f"  Mean: {df[col].mean():.2f}\n"
                    text += f"  Median: {df[col].median():.2f}\n"
                    text += f"  Min: {df[col].min():.2f}\n"
                    text += f"  Max: {df[col].max():.2f}\n"
            
            # First few rows as preview
            text += "\nData Preview:\n"
            text += df.head(10).to_string() + "\n"
            
            text += "\n"
        
        return text
    
    def _extract_text_from_csv(self, file_path: str) -> str:
        """
        Extract data from a CSV file

        Args:
            file_path: Path to the CSV file

        Returns:
            Extracted data as text
        """
        # Read CSV file
        try:
            df = pd.read_csv(file_path)
            text = "--- CSV Data ---\n"
            
            # Get column headers
            headers = df.columns.tolist()
            text += "Columns: " + ", ".join([str(h) for h in headers]) + "\n\n"
            
            # Basic statistics
            text += "Basic Statistics:\n"
            for col in df.columns:
                if pd.api.types.is_numeric_dtype(df[col]):
                    text += f"- {col}:\n"
                    text += f"  Mean: {df[col].mean():.2f}\n"
                    text += f"  Median: {df[col].median():.2f}\n"
                    text += f"  Min: {df[col].min():.2f}\n"
                    text += f"  Max: {df[col].max():.2f}\n"
            
            # First few rows as preview
            text += "\nData Preview:\n"
            text += df.head(10).to_string() + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing CSV file: {str(e)}")
            return f"ERROR: Could not process CSV file: {str(e)}"
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from a text file

        Args:
            file_path: Path to the text file

        Returns:
            Extracted text
        """
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _analyze_image(self, file_path: str) -> str:
        """
        Analyze an image using the vision model

        Args:
            file_path: Path to the image file

        Returns:
            Analysis text
        """
        if not self.use_vision_model or self.vision_model is None:
            return "Vision model is not available for image analysis."
        
        try:
            # Open image
            image = Image.open(file_path)
            
            # Generate prompt
            prompt = "Describe this image in detail, focusing on business-relevant information."
            
            # Process image
            inputs = self.vision_processor(text=prompt, images=image, return_tensors="pt").to(self.vision_model.device)
            
            # Generate
            import torch
            try:
                with torch.no_grad():
                    outputs = self.vision_model.generate(
                        **inputs,
                        max_new_tokens=500,
                        do_sample=False
                    )
                
                    # Decode within the torch.no_grad() block
                    generated_text = self.vision_processor.batch_decode(outputs, skip_special_tokens=True)[0]
            except Exception as e:
                logger.error(f"Error generating vision model output: {str(e)}")
                return f"ERROR: Failed to analyze image with vision model: {str(e)}"
            
            # Remove prompt from response
            result = generated_text.replace(prompt, "").strip()
            
            return f"IMAGE ANALYSIS:\n{result}"
            
        except Exception as e:
            logger.error(f"Error analyzing image: {str(e)}")
            return f"ERROR: Could not analyze image: {str(e)}"
    
    def _generate_gguf(self, prompt: str) -> str:
        """
        Generate text using a GGUF model

        Args:
            prompt: Prompt text

        Returns:
            Generated text
        """
        try:
            # Define generation parameters
            params = {
                "max_tokens": 2048,
                "temperature": self.temperature,
                "top_p": self.top_p,
                "stop": ["User:", "\n\nUser:"]
            }
            
            # Generate
            result = self.model(prompt, **params)
            
            # Extract and clean response
            response = result["choices"][0]["text"].strip()
            
            # Remove any leading "Assistant:" if present
            response = re.sub(r"^Assistant:\s*", "", response)
            
            return response
        except Exception as e:
            logger.error(f"Error generating response with GGUF model: {str(e)}")
            raise
    
    def _generate_hf(self, prompt: str) -> str:
        """
        Generate text using a Hugging Face model

        Args:
            prompt: Prompt text

        Returns:
            Generated text
        """
        import torch
        
        try:
            # Tokenize
            inputs = self.tokenizer(prompt, return_tensors="pt").to(self.model.device)
            
            # Generate
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs.input_ids,
                    max_new_tokens=2048,
                    temperature=self.temperature,
                    top_p=self.top_p,
                    do_sample=True,
                    pad_token_id=self.tokenizer.eos_token_id,
                )
            
            # Decode
            full_response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # Extract only the new content (after the prompt)
            response = full_response[len(prompt):].strip()
            
            # Remove any leading "Assistant:" if present
            response = re.sub(r"^Assistant:\s*", "", response)
            
            return response
        except Exception as e:
            logger.error(f"Error generating response with HF model: {str(e)}")
            raise
    
    def _extract_plot_code(self, response: str) -> List[str]:
        """
        Extract Python code that generates plots from the response

        Args:
            response: Response text

        Returns:
            List of code blocks that generate plots
        """
        # Look for Python code blocks with matplotlib or plotting functions
        plot_code_blocks = []
        
        # Find all code blocks
        code_blocks = re.findall(r'```python(.*?)```', response, re.DOTALL)
        
        # Check each code block for plotting functions
        for code in code_blocks:
            if ('plt.' in code or 'matplotlib' in code or 'sns.' in code or '.plot(' in code) and not 'plt.savefig' in code:
                # Add plt.savefig to save the plot
                if not code.strip().endswith('\n'):
                    code += '\n'
                code += '\nplt.savefig("temp_plot.png")\nplt.close()'
                plot_code_blocks.append(code)
        
        return plot_code_blocks
    
    def _execute_plot_code(self, code_blocks: List[str]) -> List[str]:
        """
        Execute code blocks to generate plots

        Args:
            code_blocks: List of code blocks

        Returns:
            List of base64-encoded plot images
        """
        plot_images = []
        
        for i, code in enumerate(code_blocks):
            # Create a temporary directory to save plots
            with tempfile.TemporaryDirectory() as temp_dir:
                plot_file = os.path.join(temp_dir, "temp_plot.png")
                
                # Create a new Python file with the code
                temp_code_file = os.path.join(temp_dir, f"plot_code_{i}.py")
                with open(temp_code_file, 'w') as f:
                    # Add necessary imports if not present
                    imports = ""
                    if 'import matplotlib.pyplot as plt' not in code:
                        imports += "import matplotlib.pyplot as plt\n"
                    if 'import numpy as np' not in code and 'np.' in code:
                        imports += "import numpy as np\n"
                    if 'import pandas as pd' not in code and 'pd.' in code:
                        imports += "import pandas as pd\n"
                    if 'import seaborn as sns' not in code and 'sns.' in code:
                        imports += "import seaborn as sns\n"
                    
                    f.write(imports + code)
                
                # Execute the code
                try:
                    exec_globals = {}
                    with open(temp_code_file, 'r') as f:
                        exec(f.read(), exec_globals)
                    
                    # Check if the plot was generated
                    if os.path.exists(plot_file):
                        # Read the image and encode as base64
                        with open(plot_file, 'rb') as f:
                            encoded_image = base64.b64encode(f.read()).decode('utf-8')
                            plot_images.append(encoded_image)
                except Exception as e:
                    logger.error(f"Error executing plot code: {str(e)}")
        
        return plot_images
    
    def _add_plot_images_to_response(self, response: str, plot_images: List[str]) -> str:
        """
        Add plot images to the response

        Args:
            response: Original response text
            plot_images: List of base64-encoded images

        Returns:
            Updated response with images
        """
        # Replace code blocks with code + image
        code_blocks = re.findall(r'```python(.*?)```', response, re.DOTALL)
        
        updated_response = response
        
        for i, code in enumerate(code_blocks):
            if i < len(plot_images):
                # Create image markdown
                image_markdown = f'\n\n![Generated Plot {i+1}](data:image/png;base64,{plot_images[i]})\n\n'
                
                # Replace code block with code block + image
                code_block = f'```python{code}```'
                replacement = f'{code_block}{image_markdown}'
                
                updated_response = updated_response.replace(code_block, replacement, 1)
        
        return updated_response
    
    def analyze_document(self, document_path: str) -> str:
        """
        Public method to analyze a document

        Args:
            document_path: Path to the document

        Returns:
            Analysis result
        """
        return self._process_document(document_path)
    
    def analyze_data(self, data: Union[str, pd.DataFrame], data_type: str = "csv") -> str:
        """
        Analyze data provided as string or DataFrame

        Args:
            data: Data to analyze (string or DataFrame)
            data_type: Type of data (csv, json, etc.)

        Returns:
            Analysis result
        """
        try:
            # Convert string data to DataFrame if needed
            if isinstance(data, str):
                if data_type.lower() == "csv":
                    df = pd.read_csv(pd.StringIO(data))
                elif data_type.lower() == "json":
                    df = pd.read_json(data)
                else:
                    return f"Unsupported data_type: {data_type}"
            else:
                df = data
            
            # Generate analysis
            analysis = "--- Data Analysis ---\n"
            
            # Basic info
            analysis += f"Dimensions: {df.shape[0]} rows x {df.shape[1]} columns\n"
            analysis += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            
            # Data types
            analysis += "Data Types:\n"
            for col, dtype in df.dtypes.items():
                analysis += f"- {col}: {dtype}\n"
            analysis += "\n"
            
            # Basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) > 0:
                analysis += "Numeric Column Statistics:\n"
                for col in numeric_cols:
                    analysis += f"- {col}:\n"
                    analysis += f"  Mean: {df[col].mean():.2f}\n"
                    analysis += f"  Median: {df[col].median():.2f}\n"
                    analysis += f"  Std Dev: {df[col].std():.2f}\n"
                    analysis += f"  Min: {df[col].min():.2f}\n"
                    analysis += f"  Max: {df[col].max():.2f}\n"
                analysis += "\n"
            
            # Missing values
            missing = df.isnull().sum()
            if missing.sum() > 0:
                analysis += "Missing Values:\n"
                for col, count in missing.items():
                    if count > 0:
                        analysis += f"- {col}: {count} ({count/df.shape[0]*100:.1f}%)\n"
                analysis += "\n"
            
            # Categorical columns
            cat_cols = df.select_dtypes(include=['object', 'category']).columns
            if len(cat_cols) > 0:
                analysis += "Categorical Variables:\n"
                for col in cat_cols:
                    value_counts = df[col].value_counts().head(5)
                    analysis += f"- {col} (top 5):\n"
                    for val, count in value_counts.items():
                        analysis += f"  {val}: {count} ({count/df.shape[0]*100:.1f}%)\n"
                analysis += "\n"
            
            # Data preview
            analysis += "Data Preview (first 5 rows):\n"
            analysis += df.head(5).to_string() + "\n"
            
            return analysis
